import pymysql
from docx import Document

# Kết nối database
conn = pymysql.connect(host='localhost', user='root', password='', db='fruit_manage', charset='utf8')
cur = conn.cursor()

# Danh sách table cần export
tables = [
    'farmers_farmer',
    'inventory_inventorystock',
    'inventory_warehouse',
    'orders_order',
    'orders_orderdetail',
    'payments_payment',
    'products_category',
    'products_product',
]

doc = Document()
doc.add_heading('Database Schema', level=1)

for tbl in tables:
    doc.add_heading(tbl, level=2)
    # Lấy column info
    cur.execute(f"SHOW COLUMNS FROM `{tbl}`")
    cols = cur.fetchall()  # [(Field, Type, Null, Key, Default, Extra), ...]

    # Tạo table trong Word (1 header + len(cols) rows)
    table = doc.add_table(rows=1 + len(cols), cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'STT', 'Tên thuộc tính', 'Kiểu dữ liệu'

    for i, col in enumerate(cols, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        row_cells[1].text = col[0]
        row_cells[2].text = col[1]

    doc.add_paragraph('')  # dòng trống giữa các bảng

# Lưu file
doc.save('db_schema.docx')
